# app_spf_Portal.py — login restored + sidebar layout + fresh DB handling
# ---------------------------------------------------------------------------------
# Restores:
#   • streamlit-authenticator login (user-specific Location access via config)
#   • All nav + controls in the **sidebar** (not top)
#   • "Data last updated" (file mtime) + "Recency Check" in sidebar
#   • Clear cache + Force re-download from GitHub + manual local DB path override
#   • Cache keys include DB signature + a manual nonce for guaranteed refresh
#
# Keep your previous app_config.yaml / Streamlit secrets layout:
#   [app_config]
#   settings.db_path = "C:/path/to/maintainx_po.db"
#   settings.quotes_db_path = "C:/path/to/quotes.db"
#   access.admin_usernames = ["brad"]
#   access.user_companies.brad = ["110 - Deckers Creek Limestone","300 - Greer Lime - Stone"] # or ["*"]
#   cookie.name = "spf_po_portal"
#   cookie.key = "change_me"
#   cookie.expiry_days = 7
#   credentials.usernames.brad.name = "Brad"
#   credentials.usernames.brad.email = "brad@example.com"
#   credentials.usernames.brad.password = "$2b$12$hashed_bcrypt_here"
#
# Optional GitHub (for "Re-download"):
#   [github]
#   repo = "YOURUSER/YOURREPO"
#   path = "maintainx_po.db"
#   branch = "main"
#   token = "ghp_..."     # only if private
# ---------------------------------------------------------------------------------

from __future__ import annotations

import os, io, re, json, sqlite3, hashlib, tempfile, time
from pathlib import Path
from typing import Optional, Tuple, List, Dict
from datetime import datetime, timezone

import pandas as pd
import streamlit as st

APP_VERSION = "2025.10.28-login+sidebar-fix"

# ---------------- Optional Word export ----------------
DOCX_OK = True
try:
    from docx import Document
    from docx.shared import Pt, Inches
except Exception:
    DOCX_OK = False

# ---------------- Auth (streamlit-authenticator) ----------------
AUTH_OK = True
try:
    import yaml
    import streamlit_authenticator as stauth
except Exception:
    AUTH_OK = False

st.set_page_config(page_title="SPF PO Portal", page_icon="📦", layout="wide")

DEFAULT_DB = "maintainx_po.db"
QUOTES_DB  = "quotes.db"
HERE = Path(__file__).resolve().parent

# ========================= Utilities =========================

def _filesig(p: Path) -> int:
    """Compact signature from mtime_ns and file size."""
    try:
        stt = p.stat()
        return (int(stt.st_mtime_ns) ^ (stt.st_size << 13)) & 0xFFFFFFFFFFFF
    except Exception:
        return 0

def _db_sig(path: str) -> int:
    try:
        return Path(path).stat().st_mtime_ns
    except Exception:
        return 0

@st.cache_data(show_spinner=False)
def q_cached(sql: str, params: Tuple, db_path: str, db_sig: int, nonce: int) -> pd.DataFrame:
    with sqlite3.connect(db_path, timeout=30, check_same_thread=False) as conn:
        try:
            conn.execute("PRAGMA journal_mode=WAL;")
            conn.execute("PRAGMA synchronous=NORMAL;")
            conn.execute("PRAGMA busy_timeout=5000;")
            conn.execute("PRAGMA foreign_keys=ON;")
        except Exception:
            pass
        return pd.read_sql_query(sql, conn, params=params)

def q(sql: str, params: tuple = (), db_path: Optional[str] = None) -> pd.DataFrame:
    path = db_path or st.session_state.get("DATA_DB_PATH", "")
    return q_cached(sql, tuple(params), path, _db_sig(path), st.session_state.get("CACHE_NONCE", 0))

@st.cache_data(show_spinner=False)
def table_columns_in_order_cached(db_path: str, table: str, db_sig: int, nonce: int) -> list[str]:
    with sqlite3.connect(db_path, timeout=30, check_same_thread=False) as conn:
        rows = conn.execute(f"PRAGMA table_info('{table}')").fetchall()
    return [r[1] for r in rows]

def table_columns_in_order(db_path: Optional[str], table: str) -> list[str]:
    path = db_path or st.session_state.get("DATA_DB_PATH", "")
    return table_columns_in_order_cached(path, table, _db_sig(path), st.session_state.get("CACHE_NONCE", 0))

def attach_row_key(df_in: pd.DataFrame) -> pd.DataFrame:
    df_in = df_in.copy()
    key_col = next((c for c in ["ID","id","Purchase Order ID","Row ID","RowID"] if c in df_in.columns), None)
    if key_col:
        df_in["__KEY__"] = df_in[key_col].astype(str); return df_in
    cols = [c for c in ["Part Number","Part Numbers","Name","Description","Vendor","Company","Created On"] if c in df_in.columns] or list(df_in.columns)
    s = df_in[cols].astype(str).agg("|".join, axis=1)
    df_in["__KEY__"] = s.apply(lambda x: hashlib.sha1(x.encode("utf-8")).hexdigest())
    return df_in

def to_xlsx_bytes(df: pd.DataFrame, sheet: str) -> bytes:
    import xlsxwriter
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine='xlsxwriter') as w:
        df.to_excel(w, index=False, sheet_name=sheet)
        ws = w.sheets[sheet]
        ws.autofilter(0,0,max(0,len(df)), max(0,len(df.columns)-1))
        for i,_ in enumerate(df.columns):
            try:
                width = 16 if df.empty else min(60, max(12, int(df.iloc[:,i].astype(str).str.len().quantile(0.9))+2))
                ws.set_column(i,i,width)
            except Exception:
                pass
    return buf.getvalue()

def build_quote_docx(company: str, date_str: str, quote_number: str,
                     vendor_text: str, ship_to_text: str, bill_to_text: str,
                     lines_df: pd.DataFrame) -> bytes:
    if not DOCX_OK:
        raise RuntimeError("python-docx not installed.")
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    hdr = doc.add_paragraph(); r = hdr.add_run(company); r.bold = True; r.font.size = Pt(14)
    title = doc.add_paragraph(); r2 = title.add_run("Quote Request"); r2.bold = True; r2.font.size = Pt(16)
    doc.add_paragraph(date_str); doc.add_paragraph(f"Quote #: {quote_number}")

    doc.add_paragraph(""); vr = doc.add_paragraph(); vr.add_run("Vendor").bold = True
    doc.add_paragraph(vendor_text if str(vendor_text).strip() else "_____________________________")

    doc.add_paragraph("")
    tbl_addr = doc.add_table(rows=2, cols=2)
    tbl_addr.rows[0].cells[0].text = "Ship To Address"
    tbl_addr.rows[0].cells[1].text = "Bill To Address"
    tbl_addr.rows[1].cells[0].text = ship_to_text
    tbl_addr.rows[1].cells[1].text = bill_to_text

    cols = ["Part Number","Description","Quantity","Price/Unit","Total"]
    lines = lines_df.copy()
    for c in cols:
        if c not in lines.columns: lines[c] = ""
    BLANK_ROWS = max(10, 30 - len(lines))
    if BLANK_ROWS>0: lines = pd.concat([lines, pd.DataFrame([dict(zip(cols, [""]*5)) for _ in range(BLANK_ROWS)])], ignore_index=True)

    doc.add_paragraph("")
    tbl = doc.add_table(rows=1 + len(lines), cols=len(cols)); tbl.style = 'Table Grid'
    widths = [Inches(1.7), Inches(3.6), Inches(0.9), Inches(1.2), Inches(1.2)]
    for j in range(len(cols)):
        for row in tbl.rows:
            row.cells[j].width = widths[j]
    for j,c in enumerate(cols): tbl.cell(0,j).text = c
    for i,(_,row) in enumerate(lines.iterrows(), start=1):
        for j,c in enumerate(cols): tbl.cell(i,j).text = str("" if pd.isna(row[c]) else row[c])

    bio = io.BytesIO(); doc.save(bio); bio.seek(0); return bio.getvalue()

# ========================= Config + Source =========================

def cfg_plain(obj):
    if isinstance(obj, dict): return {k: cfg_plain(v) for k,v in obj.items()}
    if isinstance(obj, (list,tuple)): return [cfg_plain(x) for x in obj]
    return obj

def load_config() -> dict:
    # secrets first
    if hasattr(st, "secrets") and "app_config" in st.secrets:
        return cfg_plain(st.secrets["app_config"])
    # yaml fallback
    cfg_file = HERE / "app_config.yaml"
    if cfg_file.exists():
        try:
            return yaml.safe_load(cfg_file.read_text(encoding="utf-8")) or {}
        except Exception as e:
            st.error(f"Invalid YAML in app_config.yaml: {e}")
    # default minimal
    return {
        "settings": {"db_path": str((HERE/DEFAULT_DB).resolve()),
                     "quotes_db_path": str((HERE/QUOTES_DB).resolve())},
        "access": {"admin_usernames": ["demo"], "user_companies": {"demo": ["*"]}},
        "cookie": {"name": "spf_po_portal", "key": "change_me", "expiry_days": 7},
        "credentials": {"usernames": {"demo": {"name": "Demo", "email": "demo@example.com", "password": ""}}}
    }

def get_github_cfg() -> Optional[Dict[str,str]]:
    gh = getattr(st, "secrets", {}).get("github") if hasattr(st, "secrets") else None
    if isinstance(gh, dict) and gh.get("repo") and gh.get("path"):
        return {"repo": gh.get("repo"), "path": gh.get("path"),
                "branch": gh.get("branch","main"), "token": gh.get("token")}
    return None

def download_db_from_github(repo: str, path: str, branch: str='main', token: Optional[str]=None) -> str:
    import requests
    url = f"https://api.github.com/repos/{repo}/contents/{path}?ref={branch}"
    headers = {"Accept": "application/vnd.github.v3.raw"}
    if token: headers["Authorization"] = f"token {token}"
    r = requests.get(url, headers=headers, timeout=30)
    if r.status_code != 200:
        raise RuntimeError(f"GitHub API returned {r.status_code}: {r.text[:200]}")
    tmpdir = Path(tempfile.gettempdir()) / "spf_po_cache"
    tmpdir.mkdir(parents=True, exist_ok=True)
    # write to a NEW, timestamped path to avoid stale cache reuse
    out = tmpdir / f"{int(time.time())}_{DEFAULT_DB}"
    out.write_bytes(r.content)
    # touch mtime
    now = time.time()
    os.utime(out, (now, now))
    return str(out.resolve())

def resolve_db_path(cfg: dict) -> str:
    # manual override from UI
    if st.session_state.get("MANUAL_DB_PATH") and Path(st.session_state["MANUAL_DB_PATH"]).exists():
        return st.session_state["MANUAL_DB_PATH"]

    # prefer GitHub if configured
    gh = get_github_cfg()
    if gh:
        try:
            p = download_db_from_github(**gh)
            st.session_state["GITHUB_DB_PATH"] = p
            return p
        except Exception as e:
            st.warning(f"GitHub fetch failed, falling back to local/YAML. Error: {e}")

    # YAML / local path
    yaml_db = (cfg or {}).get("settings",{}).get("db_path")
    if yaml_db:
        return str(Path(yaml_db).expanduser().resolve())

    # env var
    env_db = os.environ.get("SPF_DB_PATH")
    if env_db:
        return str(Path(env_db).expanduser().resolve())

    # default
    return str((HERE/DEFAULT_DB).resolve())

def resolve_quotes_db_path(cfg: dict) -> str:
    yaml_q = (cfg or {}).get('settings', {}).get('quotes_db_path')
    if yaml_q: return str(Path(yaml_q).expanduser().resolve())
    env_q = os.environ.get('SPF_QUOTES_DB_PATH')
    if env_q: return str(Path(env_q).expanduser().resolve())
    return str((HERE / QUOTES_DB).resolve())

# ========================= Boot + Auth =========================

cfg = load_config()

if "CACHE_NONCE" not in st.session_state: st.session_state["CACHE_NONCE"] = 0
if "DATA_DB_PATH" not in st.session_state: st.session_state["DATA_DB_PATH"] = resolve_db_path(cfg)
if "QUOTES_DB_PATH" not in st.session_state: st.session_state["QUOTES_DB_PATH"] = resolve_quotes_db_path(cfg)

DATA_DB_PATH   = st.session_state["DATA_DB_PATH"]
QUOTES_DB_PATH = st.session_state["QUOTES_DB_PATH"]

# ---- Auth UI ----
if AUTH_OK:
    cookie_cfg = cfg.get('cookie', {})
    authenticator = stauth.Authenticate(
        cfg.get('credentials', {}),
        cookie_cfg.get('name', 'spf_po_portal'),
        cookie_cfg.get('key',  'change_me'),
        cookie_cfg.get('expiry_days', 7),
    )
    name, auth_status, username = authenticator.login("Login", "main")
else:
    name, auth_status, username = ("Demo", True, "demo")

if auth_status is False:
    st.error("Username/password is incorrect")
    st.stop()
elif auth_status is None:
    st.info("Please log in.")
    st.stop()

# ---- Sidebar (layout restored here) ----
if AUTH_OK:
    authenticator.logout('Logout', 'sidebar')
st.sidebar.success(f"Logged in as {name}")

def db_info_caption(path: str) -> str:
    try:
        p = Path(path)
        ts = datetime.fromtimestamp(p.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S")
        return f"**DB:** `{p}`  \n**Last modified:** {ts}"
    except Exception:
        return f"**DB:** `{path}`"

st.sidebar.markdown("### Data Sources")
st.sidebar.caption(db_info_caption(DATA_DB_PATH))
st.sidebar.caption(f"Quotes DB: `{Path(QUOTES_DB_PATH).resolve()}`")

c1, c2 = st.sidebar.columns(2)
if c1.button("🔄 Clear cache", use_container_width=True, key="btn_clear_cache"):
    st.cache_data.clear()
    st.session_state["CACHE_NONCE"] += 1
    st.rerun()
if c2.button("🕸️ Re-download", use_container_width=True, key="btn_redownload"):
    gh = get_github_cfg()
    if gh:
        try:
            p = download_db_from_github(**gh)
            st.session_state["MANUAL_DB_PATH"] = None
            st.session_state["GITHUB_DB_PATH"] = p
            st.session_state["DATA_DB_PATH"] = p
            st.cache_data.clear()
            st.session_state["CACHE_NONCE"] += 1
            st.success("Fetched latest DB from GitHub.")
            st.rerun()
        except Exception as e:
            st.sidebar.error(f"Download failed: {e}")
    else:
        st.sidebar.warning("GitHub secrets not configured.")

st.sidebar.markdown("---")
manual = st.sidebar.text_input("Local DB path (absolute)", value="", key="manual_db_path")
if st.sidebar.button("Use this DB", use_container_width=True, key="btn_use_local"):
    if manual and Path(manual).exists():
        st.session_state["MANUAL_DB_PATH"] = str(Path(manual).resolve())
        st.session_state["DATA_DB_PATH"]   = st.session_state["MANUAL_DB_PATH"]
        st.cache_data.clear()
        st.session_state["CACHE_NONCE"] += 1
        st.rerun()
    else:
        st.sidebar.error("Path not found.")

# Recency check (sidebar)
def max_date_from(df: pd.DataFrame, candidates: List[str]) -> Optional[str]:
    for c in candidates:
        if c in df.columns:
            s = pd.to_datetime(df[c], errors="coerce")
            if s.notna().any():
                return s.max().strftime("%Y-%m-%d")
    return None

rec_cols = {
    "restock": ["Part Updated on","Last Updated","Posting Date","Approved On","Completed On","Created On","Needed By"],
    "po_outstanding": ["Created On","Approved On","Completed On","Posting Date"],
}
with st.sidebar.expander("Recency Check (max dates)"):
    for table, cols in rec_cols.items():
        try:
            df = q(f"SELECT * FROM [{table}]", db_path=DATA_DB_PATH)
            when = max_date_from(df, cols)
            st.write(f"- **{table}**: {when or '(no date columns)'}")
        except Exception as e:
            st.write(f"- **{table}**: not available ({e})")

st.sidebar.markdown("---")
st.sidebar.caption(f"Version {APP_VERSION}")

# ========================= Access control: Locations =========================

# Collect companies from data
def all_companies() -> List[str]:
    try:
        df = q("SELECT DISTINCT [Company] FROM [restock] WHERE [Company] IS NOT NULL ORDER BY 1", db_path=DATA_DB_PATH)
        return [str(x) for x in df["Company"].dropna().tolist()]
    except Exception:
        return []

all_companies_list = all_companies()

# Map allowed companies by user
username_ci = str(username).casefold()
admin_users_ci = {str(u).casefold() for u in (cfg.get('access', {}).get('admin_usernames', []) or [])}
is_admin = username_ci in admin_users_ci
uc_raw = (cfg.get('access', {}).get('user_companies', {}) or {})
uc_ci_map = {str(k).casefold(): v for k, v in uc_raw.items()}
allowed_cfg = uc_ci_map.get(username_ci, [])
if isinstance(allowed_cfg, str): allowed_cfg = [allowed_cfg]
allowed_cfg = [a for a in (allowed_cfg or [])]

def norm(s: str) -> str: return " ".join(str(s).strip().split()).casefold()
db_map = {norm(c): c for c in all_companies_list}
allowed_norm = {norm(a) for a in allowed_cfg}
star_granted = any(str(a).strip() == "*" for a in allowed_cfg)

if is_admin or star_granted:
    allowed_set = set(all_companies_list or [])
else:
    matches = {db_map[n] for n in allowed_norm if n in db_map}
    allowed_set = matches if matches else set(allowed_cfg)

if not allowed_set:
    st.error("No companies configured for your account. Ask an admin to update your access.")
    with st.expander("Company values present in data"):
        st.write(sorted(all_companies_list))
    st.stop()

company_options = sorted(allowed_set)
ADMIN_ALL = "« All companies (admin) »"
select_options = ["— Choose company —"]
if is_admin and len(all_companies_list) > 1: select_options += [ADMIN_ALL]
select_options += company_options

chosen_company = st.sidebar.selectbox("Choose your Location", options=select_options, index=0, key="company_select")

if chosen_company == "— Choose company —":
    st.info("Select your Location from the left sidebar.")
    st.stop()

if is_admin and chosen_company == ADMIN_ALL:
    chosen_companies = sorted(all_companies_list)
    title_companies = "All companies (admin)"
else:
    chosen_companies = [chosen_company]
    title_companies = chosen_company

# ========================= Page nav (sidebar) =========================

page = st.sidebar.radio("Page", ["RE-STOCK", "Outstanding POs", "Quotes"], index=0, key="page_radio")

# ========================= RE-STOCK =========================

def restock_page():
    st.markdown(f"### RE-STOCK — {title_companies}")
    try:
        df_all = q("SELECT * FROM [restock]", db_path=DATA_DB_PATH)
    except Exception as e:
        st.error(f"Could not read [restock]: {e}")
        return

    if df_all.empty:
        st.info("No RE-STOCK data found.")
        return

    df = df_all.copy()
    if "Company" in df.columns and chosen_companies:
        df = df[df["Company"].astype(str).isin([str(x) for x in chosen_companies])]

    # Search in sidebar to keep top area clean
    s = st.sidebar.text_input("Search PN / Name / Vendor contains", key="restock_search")
    if s:
        cols = [c for c in ["Part Numbers","Name","Vendor","Vendors"] if c in df.columns]
        if cols:
            m = pd.Series(False, index=df.index)
            for c in cols:
                m |= df[c].astype(str).str.contains(s, case=False, na=False)
            df = df[m]

    df = attach_row_key(df)
    df_disp = df[[c for c in df.columns if c not in {"__KEY__","ID","id","Purchase Order ID"}]].copy()
    st.dataframe(df_disp, use_container_width=True, hide_index=True, key="restock_table")

    st.download_button(
        "⬇️ Download view (.xlsx)",
        data=to_xlsx_bytes(df_disp, sheet="RE_STOCK"),
        file_name=f"RE_STOCK_{(title_companies or 'all').replace(' ','_')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key="dl_restock_xlsx"
    )

# ========================= Outstanding POs =========================

def po_outstanding_page():
    st.markdown(f"### Outstanding POs — {title_companies}")
    try:
        df_all = q("SELECT * FROM [po_outstanding]", db_path=DATA_DB_PATH)
    except Exception as e:
        st.error(f"Could not read [po_outstanding]: {e}")
        return

    if df_all.empty:
        st.info("No Outstanding POs data found."); return

    df = df_all.copy()
    if "Company" in df.columns and chosen_companies:
        df = df[df["Company"].astype(str).isin([str(x) for x in chosen_companies])]

    s = st.sidebar.text_input("Search PO # / Vendor / Part / Line contains", key="po_search")
    if s:
        cols = [c for c in ["Purchase Order #","Vendor","Part Number","Line Name"] if c in df.columns]
        if cols:
            m = pd.Series(False, index=df.index)
            for c in cols:
                m |= df[c].astype(str).str.contains(s, case=False, na=False)
            df = df[m]

    df = attach_row_key(df)
    df_disp = df[[c for c in df.columns if c not in {"__KEY__","ID","id","Purchase Order ID"}]].copy()
    st.dataframe(df_disp, use_container_width=True, hide_index=True, key="po_table")

    st.download_button(
        "⬇️ Download view (.xlsx)",
        data=to_xlsx_bytes(df_disp, sheet="Outstanding_POs"),
        file_name=f"Outstanding_POs_{(title_companies or 'all').replace(' ','_')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key="dl_po_xlsx"
    )

# ========================= Quotes (simplified) =========================

def next_quote_number(db_path: str, date_obj: datetime) -> str:
    with sqlite3.connect(db_path, timeout=30, check_same_thread=False) as conn:
        try:
            rows = conn.execute("SELECT quote_number FROM quotes WHERE quote_number LIKE ?", (f"QR-{date_obj:%Y}-%",)).fetchall()
        except Exception:
            rows = []
    used = set()
    for (qn,) in rows:
        try:
            parts = str(qn).split("-"); seq = int(parts[-1])
            if str(parts[1]) == f"{date_obj:%Y}": used.add(seq)
        except Exception:
            pass
    seq = 1
    while seq in used: seq += 1
    return f"QR-{date_obj:%Y}-{seq:04d}"

def ensure_quotes_table(db_path: str) -> None:
    with sqlite3.connect(db_path, timeout=30, check_same_thread=False) as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS quotes (
              id INTEGER PRIMARY KEY AUTOINCREMENT,
              quote_number TEXT UNIQUE,
              company TEXT,
              created_by TEXT,
              vendor TEXT,
              ship_to TEXT,
              bill_to TEXT,
              quote_date TEXT,
              status TEXT,
              source TEXT,
              lines_json TEXT,
              updated_at TEXT
            )
        """); conn.commit()

def quotes_page():
    ensure_quotes_table(QUOTES_DB_PATH)
    st.markdown(f"### Quotes — {title_companies}")

    # New Quote
    st.subheader("New Quote")
    c1, c2 = st.columns([1,1])
    with c1:
        company_new = st.selectbox("Location", options=(company_options or [""]), index=0, key="new_quote_location")
        vendor      = st.text_input("Vendor", value="", key="new_quote_vendor")
        quote_no    = st.text_input("Quote #", value=next_quote_number(QUOTES_DB_PATH, datetime.utcnow()), key="new_quote_no")
    with c2:
        ship_to = st.text_area("Ship To", value="", height=120, key="new_quote_ship")
        bill_to = st.text_area("Bill To", value="", height=120, key="new_quote_bill")

    lines = st.data_editor(
        pd.DataFrame([{"Part Number":"","Description":"","Quantity":"","Price/Unit":"","Total":""} for _ in range(12)]),
        hide_index=True, use_container_width=True, num_rows="dynamic",
        column_config={
            "Part Number": st.column_config.TextColumn("Part Number"),
            "Description": st.column_config.TextColumn("Description"),
            "Quantity":    st.column_config.NumberColumn("Qty", min_value=0, step=1),
            "Price/Unit":  st.column_config.TextColumn("Price/Unit"),
            "Total":       st.column_config.TextColumn("Total"),
        },
        key="new_quote_editor"
    )

    a1, a2, _ = st.columns([1,1,5])
    if a1.button("Save Quote", key="btn_save_quote"):
        payload = json.dumps(lines.fillna("").astype(str).to_dict(orient="records"), ensure_ascii=False)
        now = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S")
        with sqlite3.connect(QUOTES_DB_PATH, timeout=30, check_same_thread=False) as conn:
            try:
                conn.execute("""
                    INSERT INTO quotes(quote_number, company, created_by, vendor, ship_to, bill_to,
                                       quote_date, status, source, lines_json, updated_at)
                    VALUES (?, ?, ?, ?, ?, ?, date('now'), 'draft', 'manual', ?, ?)
                """, (quote_no, company_new, str(st.session_state.get("username", "")), vendor, ship_to, bill_to, payload, now))
                conn.commit()
                st.success(f"Saved quote {quote_no}")
            except sqlite3.IntegrityError:
                st.error("Quote # already exists. Change number and try again.")

    if a2.button("Download Word", key="btn_dl_quote_word"):
        try:
            doc_bytes = build_quote_docx(
                company=company_new, date_str=datetime.now().strftime("%Y-%m-%d"),
                quote_number=quote_no, vendor_text=vendor,
                ship_to_text=ship_to, bill_to_text=bill_to, lines_df=lines
            )
            st.download_button(
                "Download Quote (Word)",
                data=doc_bytes,
                file_name=f"{quote_no}_{re.sub(r'[^A-Za-z0-9._ -]+','_',company_new)}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="dl_quote_word_now"
            )
        except Exception as e:
            st.error(f"Could not build Word file: {e}")

    st.divider()
    st.subheader("Saved Quotes")
    with sqlite3.connect(QUOTES_DB_PATH, timeout=30, check_same_thread=False) as conn:
        try:
            dfq = pd.read_sql_query(
                "SELECT id, quote_number, quote_date, company, vendor, status, length(lines_json) AS bytes FROM quotes ORDER BY id DESC",
                conn
            )
        except Exception:
            dfq = pd.DataFrame()
    if dfq.empty:
        st.info("No saved quotes yet.")
    else:
        st.dataframe(dfq, use_container_width=True, hide_index=True)

# ========================= Router =========================

if page == "RE-STOCK":
    restock_page()
elif page == "Outstanding POs":
    po_outstanding_page()
else:
    quotes_page()

